from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path("Bio Paper Codex Editor - working.docx")
OUT = Path("Bio Paper Codex Editor - graph info styled.docx")


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=150, start=190, bottom=160, end=190):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="000000", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9000):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))

    grid = table._tbl.tblGrid
    for col in list(grid):
        grid.remove(col)
    col = OxmlElement("w:gridCol")
    col.set(qn("w:w"), str(width_dxa))
    grid.append(col)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width_dxa))


def format_run(run, *, bold=False, size=10, color="FFFFFF"):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def format_para(paragraph, *, before=0, after=6, line=1.08):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def build_panel(doc, caption_paragraph, info_paragraphs):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_fill(cell, "000000")
    set_cell_margins(cell)

    title = cell.paragraphs[0]
    format_para(title, after=5)
    run = title.add_run("INFORMATION")
    format_run(run, bold=True, size=8, color="FFFFFF")

    for p in info_paragraphs:
        para = cell.add_paragraph()
        format_para(para, after=5)
        text = p.text.strip()
        if text.startswith("y=") or "exp(" in text:
            size = 8.5
        else:
            size = 9.5
        run = para.add_run(text)
        format_run(run, size=size, color="FFFFFF")

    # Relocate the completed table from the end of the body to immediately
    # after the caption it describes.
    tbl = table._tbl
    body = tbl.getparent()
    body.remove(tbl)
    caption_paragraph._p.addnext(tbl)


def main():
    doc = Document(SOURCE)
    paragraphs = list(doc.paragraphs)
    caption_indices = [
        i for i, p in enumerate(paragraphs)
        if p.text.strip().startswith("Figure ") and " |" in p.text
    ]

    groups = []
    for pos, idx in enumerate(caption_indices):
        end = caption_indices[pos + 1] if pos + 1 < len(caption_indices) else len(paragraphs)
        info = [p for p in paragraphs[idx + 1:end] if p.text.strip()]
        if info:
            groups.append((paragraphs[idx], info))

    for caption, info in reversed(groups):
        build_panel(doc, caption, info)
        for p in info:
            remove_paragraph(p)

    doc.save(OUT)
    print(f"Styled {len(groups)} graph information sections -> {OUT}")


if __name__ == "__main__":
    main()
